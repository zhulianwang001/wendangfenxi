import streamlit as st
import requests
import json
import pandas as pd
import datetime
import os
from io import BytesIO

# 处理文档的库
import docx2txt
import PyPDF2
import pdfplumber

# 设置页面配置
st.set_page_config(
    page_title="法律法规合规性分析工具",
    page_icon="⚖️",
    layout="wide"
)

# 应用标题和说明
st.title("法律法规合规性分析工具")
st.markdown("""
本应用用于对上传的文档(Word, PDF)进行法律法规的科学性、合理性和合规性分析，基于DeepSeek API。
分析结果将指出存在的潜在问题，并提供修改完善的意见建议。
""")

# DeepSeek API配置
st.sidebar.header("API配置")
api_key = st.sidebar.text_input("DeepSeek API密钥", type="password")
api_base = st.sidebar.text_input("API Base URL", value="https://api.deepseek.com/v1")

# 模型选择
model_options = ["deepseek-chat", "deepseek-coder"]
selected_model = st.sidebar.selectbox("选择模型", model_options)

# 文档处理函数
def extract_text_from_docx(file):
    """从Word文档中提取文本"""
    try:
        text = docx2txt.process(file)
        return text
    except Exception as e:
        return f"文档处理错误: {str(e)}"

def extract_text_from_pdf(file):
    """从PDF文档中提取文本"""
    text = ""
    try:
        # 使用pdfplumber提取文本，它对中文支持更好
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        
        # 如果pdfplumber提取的文本为空，尝试使用PyPDF2
        if not text.strip():
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        return text
    except Exception as e:
        return f"PDF处理错误: {str(e)}"

# 主页面
st.header("文档分析")

# 文件上传区域
uploaded_file = st.file_uploader("上传文档 (Word或PDF)", type=["docx", "pdf"])

# 提取的文本显示和编辑区域
extracted_text = ""
if uploaded_file is not None:
    # 根据文件类型提取文本
    file_extension = uploaded_file.name.split(".")[-1].lower()
    
    if file_extension == "docx":
        with st.spinner("正在提取Word文档内容..."):
            extracted_text = extract_text_from_docx(uploaded_file)
    elif file_extension == "pdf":
        with st.spinner("正在提取PDF文档内容..."):
            extracted_text = extract_text_from_pdf(uploaded_file)
    
    # 显示提取的文本并允许编辑
    st.subheader("提取的文本内容")
    text_to_analyze = st.text_area("可以编辑文本内容", extracted_text, height=200)
else:
    text_to_analyze = st.text_area("或直接输入文本内容", "", height=200)

# 分析类型选择
analysis_types = st.multiselect(
    "选择分析类型",
    ["科学性分析", "合理性分析", "合规性分析"],
    default=["科学性分析", "合理性分析", "合规性分析"]
)

# 关联法律法规领域
legal_areas = st.multiselect(
    "关联法律法规领域",
    ["公司法", "合同法", "劳动法", "知识产权法", "网络安全法", "个人信息保护法", "数据安全法", "反垄断法", "消费者权益保护法", "其他"],
    default=["公司法", "合同法"]
)

# 自定义分析提示词
custom_prompt = st.text_area(
    "自定义分析提示词(可选)",
    value="",
    help="您可以输入特定的分析要求，例如针对某个具体法规的分析要点"
)

# 分析函数
def analyze_text(text, analysis_types, legal_areas, custom_prompt):
    if not api_key:
        return {"error": "请提供有效的DeepSeek API密钥"}
    
    # 构建请求的分析提示词
    prompt = f"""
请作为专业的法律顾问，对以下文本内容进行详细分析:

---文本开始---
{text}
---文本结束---

分析要求:
"""
    
    # 添加分析类型
    if "科学性分析" in analysis_types:
        prompt += """
1. 科学性分析:
   - 文本内容是否基于科学事实和证据
   - 是否存在科学谬误或不准确的表述
   - 科学依据是否充分
"""
    
    if "合理性分析" in analysis_types:
        prompt += """
2. 合理性分析:
   - 文本内容的逻辑是否合理
   - 表述是否存在矛盾之处
   - 实施的可行性评估
"""
    
    if "合规性分析" in analysis_types:
        prompt += """
3. 合规性分析:
   - 文本内容是否符合中国现行法律法规
   - 具体列出潜在违反的法律法规条款
   - 合规风险等级评估(低/中/高)
"""
    
    # 添加关联法律法规领域
    if legal_areas:
        prompt += f"""
4. 重点关注以下法律法规领域:
   - {', '.join(legal_areas)}
"""
    
    # 添加自定义提示词
    if custom_prompt:
        prompt += f"""
5. 其他特定分析要求:
   {custom_prompt}
"""
    
    # 添加输出格式要求
    prompt += """
请提供分析报告，格式要求:
1. 综合评估: 总体合规性评价
2. 问题清单: 列出发现的所有问题，并标注问题严重程度(轻微/中度/严重)
3. 修改建议: 针对每个问题提供具体的修改建议
4. 法律依据: 引用相关法律法规条款

请确保分析报告客观、全面、专业，并提供切实可行的建议。
"""
    
    # 构建API请求
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    
    payload = {
        "model": selected_model,
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ],
        "temperature": 0.3,
        "max_tokens": 4000
    }
    
    try:
        response = requests.post(
            f"{api_base}/chat/completions",
            headers=headers,
            data=json.dumps(payload)
        )
        
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        else:
            return {
                "error": f"API请求失败: {response.status_code}",
                "details": response.text
            }
    except Exception as e:
        return {"error": f"请求过程中发生错误: {str(e)}"}

# 创建DOCX报告
def create_docx_report(analysis_result, input_text):
    from docx import Document
    
    doc = Document()
    
    # 添加标题
    doc.add_heading('法律法规合规性分析报告', 0)
    
    # 添加分析时间
    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f'生成时间: {current_time}')
    
    # 添加分隔线
    doc.add_paragraph('---------------------------------------------------')
    
    # 添加原始文本
    doc.add_heading('原始文本', level=1)
    doc.add_paragraph(input_text)
    
    # 添加分隔线
    doc.add_paragraph('---------------------------------------------------')
    
    # 添加分析结果
    doc.add_heading('分析结果', level=1)
    doc.add_paragraph(analysis_result)
    
    # 保存到内存中
    docx_io = BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    return docx_io

# 创建PDF报告
def create_pdf_report(analysis_result, input_text):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    
    pdf_io = BytesIO()
    doc = SimpleDocTemplate(pdf_io, pagesize=A4)
    styles = getSampleStyleSheet()
    
    # 创建自定义样式
    title_style = styles["Title"]
    heading_style = styles["Heading1"]
    normal_style = styles["Normal"]
    
    # 添加中文字体支持
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # 尝试注册Arial Unicode MS字体（通用Unicode字体，包含中文）
        try:
            pdfmetrics.registerFont(TTFont('Arial', 'Arial Unicode.ttf'))
            normal_style = ParagraphStyle('CustomNormal', parent=styles['Normal'], fontName='Arial')
            title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontName='Arial')
            heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading1'], fontName='Arial')
        except:
            # 如果Arial Unicode不可用，可以尝试其他可能有的字体
            pass
    except ImportError:
        pass
    
    # 构建报告内容
    report_content = []
    
    # 添加标题
    report_content.append(Paragraph("法律法规合规性分析报告", title_style))
    report_content.append(Spacer(1, 20))
    
    # 添加分析时间
    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    report_content.append(Paragraph(f"生成时间: {current_time}", normal_style))
    report_content.append(Spacer(1, 20))
    
    # 添加原始文本
    report_content.append(Paragraph("原始文本", heading_style))
    report_content.append(Spacer(1, 10))
    report_content.append(Paragraph(input_text, normal_style))
    report_content.append(Spacer(1, 20))
    
    # 添加分析结果
    report_content.append(Paragraph("分析结果", heading_style))
    report_content.append(Spacer(1, 10))
    
    # 将分析结果拆分为段落
    for paragraph in analysis_result.split('\n\n'):
        if paragraph.strip():
            report_content.append(Paragraph(paragraph, normal_style))
            report_content.append(Spacer(1, 10))
    
    # 构建PDF
    doc.build(report_content)
    pdf_io.seek(0)
    
    return pdf_io

# 分析按钮
if st.button("进行分析"):
    if not text_to_analyze:
        st.error("请上传文档或输入需要分析的文本内容")
    elif not api_key:
        st.error("请提供DeepSeek API密钥")
    else:
        with st.spinner("正在进行分析，请稍候..."):
            result = analyze_text(
                text_to_analyze, 
                analysis_types, 
                legal_areas, 
                custom_prompt
            )
            
            if isinstance(result, dict) and "error" in result:
                st.error(f"分析错误: {result['error']}")
                if "details" in result:
                    st.error(f"详细信息: {result['details']}")
            else:
                st.success("分析完成!")
                
                # 显示分析结果
                st.header("分析结果")
                st.markdown(result)
                
                # 保存分析结果以供下载
                st.session_state.analysis_result = result
                st.session_state.input_text = text_to_analyze

# 如果分析结果存在，提供下载选项
if 'analysis_result' in st.session_state and 'input_text' in st.session_state:
    st.header("报告下载")
    
    # 提供不同格式的下载选项
    col1, col2, col3 = st.columns(3)
    
    # TXT下载
    with col1:
        txt_data = st.session_state.analysis_result
        st.download_button(
            label="下载TXT报告",
            data=txt_data,
            file_name=f"法律分析报告_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            mime="text/plain"
        )
    
    # DOCX下载
    with col2:
        try:
            docx_data = create_docx_report(st.session_state.analysis_result, st.session_state.input_text)
            st.download_button(
                label="下载DOCX报告",
                data=docx_data,
                file_name=f"法律分析报告_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.warning(f"DOCX报告生成失败: {str(e)}")
            st.info("提示: 需要安装 python-docx 库才能生成DOCX报告")
    
    # PDF下载
    with col3:
        try:
            pdf_data = create_pdf_report(st.session_state.analysis_result, st.session_state.input_text)
            st.download_button(
                label="下载PDF报告",
                data=pdf_data,
                file_name=f"法律分析报告_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf"
            )
        except Exception as e:
            st.warning(f"PDF报告生成失败: {str(e)}")
            st.info("提示: 需要安装 reportlab 库才能生成PDF报告")

# 帮助信息
with st.expander("使用说明"):
    st.markdown("""
    ### 使用方法
    1. 在侧边栏中输入您的DeepSeek API密钥
    2. 上传Word或PDF文档，或直接在文本框中输入内容
    3. 如有需要，可以编辑提取的文本内容
    4. 选择需要的分析类型和关联法律法规领域
    5. 点击"进行分析"按钮
    6. 分析完成后可下载不同格式的报告

    ### 依赖库安装
    请确保安装以下库：
    ```
    pip install streamlit requests pandas docx2txt PyPDF2 pdfplumber python-docx reportlab
    ```
    """)

# 页脚
st.markdown("---")
st.markdown("© 2025 法律法规合规性分析工具 | Powered by DeepSeek API")